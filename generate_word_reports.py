from pathlib import Path
import csv

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\lenovo\Documents\New project")


def load_csv_rows(csv_path: Path):
    with csv_path.open("r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def add_table(document: Document, rows):
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "迭代次数"
    hdr_cells[1].text = "w"
    hdr_cells[2].text = "b"
    hdr_cells[3].text = "Loss"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(int(float(row["Iteration"])))
        cells[1].text = f"{float(row['w']):.10f}"
        cells[2].text = f"{float(row['b']):.10f}"
        cells[3].text = f"{float(row['Loss']):.10f}"


def build_report(version_dir: Path, version_name: str, out_name: str):
    first30 = load_csv_rows(version_dir / "iteration_first30.csv")
    all50 = load_csv_rows(version_dir / "iteration_all50.csv")
    final = all50[-1]

    doc = Document()

    title = doc.add_heading(f"第一次实验课作业报告（{version_name}）", 0)
    title.alignment = 1

    p = doc.add_paragraph("实验任务：使用梯度下降法对 100 组数据进行线性拟合。")
    p.runs[0].font.size = Pt(12)

    doc.add_heading("1. 模型与目标函数", level=1)
    doc.add_paragraph("线性模型：y_hat = w * x + b")
    doc.add_paragraph("损失函数：Loss = (1/N) * Σ(y_hat - y)^2，N = 100")

    doc.add_heading("2. 参数设置", level=1)
    doc.add_paragraph("学习率：0.3")
    doc.add_paragraph("最大迭代次数：50")
    doc.add_paragraph("初始参数：w = 0.5，b = 0.5")

    doc.add_heading("3. 最终结果（第 50 次迭代）", level=1)
    doc.add_paragraph(f"w = {float(final['w']):.10f}")
    doc.add_paragraph(f"b = {float(final['b']):.10f}")
    doc.add_paragraph(f"Loss = {float(final['Loss']):.10f}")
    doc.add_paragraph(
        f"拟合直线：y = {float(final['w']):.10f} * x + {float(final['b']):.10f}"
    )

    doc.add_heading("4. 前 30 次迭代记录（w、b、Loss）", level=1)
    add_table(doc, first30)

    doc.add_heading("5. 图像结果", level=1)
    doc.add_paragraph("图 1：50 次迭代后拟合直线")
    doc.add_picture(str(version_dir / "fit_line_after_50.png"), width=Inches(6.2))
    doc.add_paragraph("图 2：Loss 随迭代次数变化曲线（50 次）")
    doc.add_picture(str(version_dir / "loss_curve_50.png"), width=Inches(6.2))

    doc.add_heading("6. 结论", level=1)
    doc.add_paragraph(
        "梯度下降在 50 次迭代内稳定收敛，Loss 持续下降，得到可用的线性拟合结果。"
    )

    out_path = version_dir / out_name
    doc.save(str(out_path))
    return out_path


def main():
    py_doc = build_report(
        ROOT / "python_version", "Python 版本", "实验报告_Python版.docx"
    )
    matlab_doc = build_report(
        ROOT / "matlab_version", "MATLAB 版本", "实验报告_MATLAB版.docx"
    )
    print(f"Generated: {py_doc}")
    print(f"Generated: {matlab_doc}")


if __name__ == "__main__":
    main()
